"""
Literature review multi-format exporter for Baṣīra.

Generates DOCX (via python-docx) and PDF (via Markdown → HTML → WeasyPrint).
"""
import io
import re
from datetime import datetime, timezone


# ── Shared Markdown builder ────────────────────────────────────────────────────


def review_to_markdown(review: dict) -> str:
    """Build a complete Markdown document from a LiteratureReview dict."""
    lines: list[str] = []
    topic = review.get("topic", "Literature Review")
    lines.append(f"# Literature review: {topic}")
    lines.append("")
    lines.append(
        f"Window: {review.get('window_days', '?')}d "
        f"· Generated: {str(review.get('created_at', ''))[:10]}"
    )
    lines.append("")

    for c in review.get("clusters", []):
        label = c.get("cluster_label", "Cluster")
        lines.append(f"## {label}")
        lines.append("")
        lines.append(c.get("synthesis", ""))
        lines.append("")

        comp = c.get("comparison_table", [])
        if comp:
            lines.append("| Work | Method | Dataset | Key result |")
            lines.append("| --- | --- | --- | --- |")
            for row in comp:
                esc = lambda s: str(s or "").replace("|", "\\|").replace("\n", " ")
                lines.append(
                    f"| {esc(row.get('work', ''))} "
                    f"| {esc(row.get('method', ''))} "
                    f"| {esc(row.get('dataset', ''))} "
                    f"| {esc(row.get('key_result', ''))} |"
                )
            lines.append("")

        gaps = c.get("gaps", [])
        if gaps:
            lines.append("### Gaps")
            for g in gaps:
                lines.append(f"- {g}")
            lines.append("")

        top = c.get("top_cite", "")
        if top:
            lines.append(f"**Top cite:** {top}")
            lines.append("")

    return "\n".join(lines)


# ── DOCX export ────────────────────────────────────────────────────────────────


def export_review_docx(review: dict) -> io.BytesIO:
    """Generate a .docx file from a literature review and return as BytesIO."""
    from docx import Document  # noqa: PLC0415
    from docx.shared import Pt, Inches, RGBColor  # noqa: PLC0415
    from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: PLC0415

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(4)

    # Title
    title = doc.add_heading(f"Literature review: {review.get('topic', '')}", level=1)
    for run in title.runs:
        run.font.size = Pt(16)

    # Metadata
    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        f"Window: {review.get('window_days', '?')}d  ·  "
        f"Generated: {str(review.get('created_at', ''))[:10]}"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for c in review.get("clusters", []):
        doc.add_heading(c.get("cluster_label", "Cluster"), level=2)
        doc.add_paragraph(c.get("synthesis", ""))

        comp = c.get("comparison_table", [])
        if comp:
            table = doc.add_table(rows=1 + len(comp), cols=4)
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            hdr = table.rows[0].cells
            for i, label in enumerate(["Work", "Method", "Dataset", "Key result"]):
                hdr[i].text = label

            for ri, row in enumerate(comp):
                cells = table.rows[ri + 1].cells
                for ci, key in enumerate(["work", "method", "dataset", "key_result"]):
                    cells[ci].text = str(row.get(key, ""))

            doc.add_paragraph()  # spacer

        gaps = c.get("gaps", [])
        if gaps:
            doc.add_heading("Gaps", level=3)
            for g in gaps:
                doc.add_paragraph(g, style="List Bullet")

        top = c.get("top_cite", "")
        if top:
            p = doc.add_paragraph()
            run = p.add_run("Top cite: ")
            run.bold = True
            p.add_run(top)

        doc.add_paragraph()  # spacer between clusters

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── PDF export ─────────────────────────────────────────────────────────────────


def export_review_pdf(review: dict) -> io.BytesIO:
    """Generate a PDF via Markdown → HTML → WeasyPrint."""
    md = review_to_markdown(review)

    html_body = _md_to_html(md)

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<style>
  body {{ font-family: 'DejaVu Serif', serif; font-size: 11pt; line-height: 1.5; color: #222; max-width: 700px; margin: 40px auto; }}
  h1 {{ font-size: 18pt; border-bottom: 2px solid #2F6FED; padding-bottom: 6px; color: #111; }}
  h2 {{ font-size: 14pt; margin-top: 24px; color: #2F6FED; }}
  h3 {{ font-size: 12pt; margin-top: 16px; color: #333; }}
  table {{ border-collapse: collapse; width: 100%; margin: 12px 0; font-size: 10pt; }}
  th, td {{ border: 1px solid #ccc; padding: 6px 8px; text-align: left; }}
  th {{ background: #f0f0f0; font-weight: 600; }}
  ul {{ margin: 6px 0; padding-left: 20px; }}
  li {{ margin-bottom: 2px; }}
  p {{ margin: 6px 0; }}
  .meta {{ color: #888; font-size: 9pt; }}
  strong {{ font-weight: 600; }}
</style>
</head>
<body>
  <p class="meta">
    Window: {review.get('window_days', '?')}d  ·  Generated: {str(review.get('created_at', ''))[:10]}
  </p>
  {html_body}
</body>
</html>"""

    from weasyprint import HTML  # noqa: PLC0415

    buf = io.BytesIO()
    HTML(string=html).write_pdf(buf)
    buf.seek(0)
    return buf


# ── Helpers ────────────────────────────────────────────────────────────────────


def _md_to_html(md_text: str) -> str:
    """Convert basic Markdown subset to HTML using the `markdown` library."""
    import markdown as md_lib  # noqa: PLC0415

    return md_lib.markdown(
        md_text,
        extensions=["extra", "sane_lists"],
    )
